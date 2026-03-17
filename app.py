from flask import Flask, request, jsonify, render_template, send_file
import zipfile
import os
import tempfile
import pdfplumber
import openpyxl
import base64
import io
import re
import uuid
from pathlib import Path

# Read .env directly — handles BOM, quotes, CRLF, and spaces
_env_path = Path(__file__).parent / '.env'
if _env_path.exists():
    with open(_env_path, encoding='utf-8-sig') as _ef:
        for _line in _ef:
            _line = _line.strip()
            if _line and not _line.startswith('#') and '=' in _line:
                _k, _, _v = _line.partition('=')
                _k = _k.strip()
                _v = _v.strip().strip('"').strip("'")
                if _k and _v:
                    os.environ.setdefault(_k, _v)
                    print(f"[.env] Loaded: {_k} = {_v[:6]}...")

try:
    from pdf2image import convert_from_bytes
    from docx import Document
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False

try:
    from google.cloud import vision
    GOOGLE_VISION_AVAILABLE = True
except ImportError:
    GOOGLE_VISION_AVAILABLE = False


app = Flask(__name__)
UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), 'bahamas_uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

_NO_TEXT_PLACEHOLDERS = frozenset([
    "[Scanned PDF - OCR returned no text]",
    "[No text could be extracted from this PDF]",
])


def _google_vision_ocr(image_bytes):
    """OCR a single image using Google Cloud Vision API."""
    if not GOOGLE_VISION_AVAILABLE:
        return ''
    try:
        client = vision.ImageAnnotatorClient()
        image = vision.Image(content=image_bytes)
        response = client.text_detection(image=image)
        if response.error.message:
            print(f"Vision API error: {response.error.message}")
            return ''
        texts = response.text_annotations
        return texts[0].description if texts else ''
    except Exception as e:
        print(f"Google Vision OCR error: {e}")
        return ''


def extract_text_with_ocr(pdf_data, max_pages=5):
    """OCR a PDF using Google Cloud Vision API."""
    if not PDF2IMAGE_AVAILABLE:
        return ''
    if not GOOGLE_VISION_AVAILABLE:
        print("Google Vision API not available - install google-cloud-vision")
        return ''
    try:
        images = convert_from_bytes(pdf_data, dpi=150, last_page=max_pages)
        page_texts = []
        for i, img in enumerate(images):
            try:
                img_buffer = io.BytesIO()
                img.save(img_buffer, format='PNG')
                img_bytes = img_buffer.getvalue()
                text = _google_vision_ocr(img_bytes)
                if text:
                    page_texts.append(text)
            except Exception as ocr_e:
                print(f"OCR page {i+1} error: {ocr_e}")
            finally:
                img.close()
        del images
        return '\n'.join(page_texts)
    except Exception as e:
        print(f"OCR ERROR: {str(e)}")
        return ''


def extract_data_from_image(image_b64, mime, filename):
    """Extract data from image using Google Vision OCR + regex."""
    if not GOOGLE_VISION_AVAILABLE:
        return None
    try:
        image_bytes = base64.b64decode(image_b64)
        text = _google_vision_ocr(image_bytes)
        if not text:
            return None

        result = {'company_info': {}}

        # Try all extractors
        general = extract_general(text)
        mgt = extract_mgt7(text, [])
        aoc = extract_aoc4(text, [])

        # Merge company_info
        for data in [general, mgt, aoc]:
            for k, v in data.get('company_info', {}).items():
                result['company_info'].setdefault(k, v)

        # Add other fields
        if mgt.get('directors'):
            result['directors'] = [
                {'din': d[0] if len(d) > 0 else '', 'name': d[1] if len(d) > 1 else '', 'designation': d[2] if len(d) > 2 else ''}
                for d in mgt['directors']
            ]
        if aoc.get('auditor'):
            result['auditor'] = aoc['auditor']

        return result if result['company_info'] else None
    except Exception as e:
        print(f"Vision OCR extraction error: {e}")
        return None


def extract_structured_from_text(text):
    """Structure text using regex patterns."""
    sections = []

    # Try to extract known fields using regex
    extracted = extract_general(text)
    if extracted.get('company_info'):
        sections.append({
            'heading': 'Company Information',
            'type': 'fields',
            'data': extracted['company_info']
        })

    # Try MGT-7 extraction
    mgt_data = extract_mgt7(text, [])
    if mgt_data.get('company_info'):
        for k, v in mgt_data['company_info'].items():
            if 'Company Information' not in [s.get('heading') for s in sections]:
                sections.append({'heading': 'Company Information', 'type': 'fields', 'data': {}})
            for s in sections:
                if s.get('heading') == 'Company Information':
                    s['data'][k] = v
    if mgt_data.get('employees'):
        sections.append({'heading': 'Employee Details', 'type': 'fields', 'data': mgt_data['employees']})
    if mgt_data.get('directors'):
        sections.append({'heading': 'Directors', 'type': 'table', 'headers': ['DIN', 'Name', 'Designation'], 'rows': mgt_data['directors']})

    # Try AOC-4 extraction
    aoc_data = extract_aoc4(text, [])
    if aoc_data.get('auditor'):
        sections.append({'heading': 'Auditor Details', 'type': 'fields', 'data': aoc_data['auditor']})

    # If no structured data found, return raw text
    if not sections:
        sections.append({'heading': 'Extracted Text', 'type': 'text', 'content': text[:5000]})

    return sections


@app.errorhandler(500)
def internal_error(err):
    return jsonify({'error': f'Internal server error: {str(err)}'}), 500

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Please upload a smaller file.'}), 413

@app.route('/health')
def health():
    return jsonify({
        'status': 'ok',
        'google_vision_available': GOOGLE_VISION_AVAILABLE,
        'pdf2image_available': PDF2IMAGE_AVAILABLE,
    })

@app.route('/')
def home():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload():
    try:
        return _upload_inner()
    except Exception as exc:
        print(f"Unhandled upload error: {exc}")
        return jsonify({'error': f'Unexpected server error: {str(exc)}'}), 500


def _upload_inner():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'})

    file = request.files['file']
    lower_name = file.filename.lower()

    # Single image upload — OCR using Google Vision API
    image_exts = ('.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.webp')
    if lower_name.endswith(image_exts):
        image_bytes = file.read()
        text = ''
        if GOOGLE_VISION_AVAILABLE:
            try:
                text = _google_vision_ocr(image_bytes).strip()
            except Exception as e:
                print(f"Image OCR error: {e}")
        if text:
            structured = [{'heading': 'Extracted Text', 'type': 'text', 'content': text}]
        else:
            structured = [{'heading': 'No Text Found', 'type': 'text',
                           'content': 'Could not extract text from this image. Check Google Vision API credentials.'}]
        return jsonify({
            'type': 'image',
            'filename': file.filename,
            'structured': structured
        })

    # Standalone PDF
    if lower_name.endswith('.pdf'):
        pdf_data = file.read()
        text = ''
        try:
            with pdfplumber.open(io.BytesIO(pdf_data)) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + '\n'
        except Exception as e:
            print(f"pdfplumber error: {e}")
        if not text.strip():
            text = extract_text_with_ocr(pdf_data)
        if not text.strip():
            text = '[No text could be extracted from this PDF]'
        if text in _NO_TEXT_PLACEHOLDERS:
            structured = [{'heading': 'No Text Extracted', 'type': 'text',
                            'content': 'Could not extract text from this PDF.'}]
        else:
            structured = extract_structured_from_text(text)
        return jsonify({'type': 'image', 'filename': file.filename, 'structured': structured})

    # Standalone Word (.docx)
    if lower_name.endswith('.docx'):
        if not PDF2IMAGE_AVAILABLE:
            return jsonify({'error': 'Word file support requires python-docx and pdf2image to be installed'})
        doc_data = file.read()
        try:
            doc = Document(io.BytesIO(doc_data))
            text = '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
        except Exception as e:
            return jsonify({'error': f'Could not read Word file: {str(e)}'})
        structured = extract_structured_from_text(text)
        return jsonify({'type': 'image', 'filename': file.filename, 'structured': structured})

    # Standalone Excel
    if lower_name.endswith(('.xlsx', '.xls', '.xlsm')):
        excel_data = file.read()
        try:
            wb = openpyxl.load_workbook(io.BytesIO(excel_data), data_only=True)
            text = ''
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text += f'Sheet: {sheet}\n'
                for row in ws.iter_rows(values_only=True):
                    row_text = ' | '.join(str(c) if c is not None else '' for c in row)
                    if row_text.strip('| '):
                        text += row_text + '\n'
        except Exception as e:
            return jsonify({'error': f'Could not read Excel file: {str(e)}'})
        structured = extract_structured_from_text(text)
        return jsonify({'type': 'image', 'filename': file.filename, 'structured': structured})

    if not lower_name.endswith('.zip'):
        return jsonify({'error': 'Unsupported file type. Please upload a PDF, image, Word (.docx), Excel, or ZIP file.'})

    safe_name = os.path.basename(file.filename)
    if not safe_name:
        return jsonify({'error': 'Invalid filename'})
    unique_name = f"{uuid.uuid4().hex}_{safe_name}"
    zip_path = os.path.join(UPLOAD_FOLDER, unique_name)
    file.save(zip_path)

    results = {
        'company_info': {},
        'directors': [],
        'shareholders': [],
        'auditor': {},
        'business_activity': [],
        'subsidiaries': [],
        'employees': {},
        'financial': {},
        'images': [],
        'raw_texts': []
    }

    if not zipfile.is_zipfile(zip_path):
        os.remove(zip_path)
        return jsonify({'error': 'Invalid or corrupted ZIP file'})

    # Read all entries up front so the ZipFile handle can be closed before processing
    entries = []
    with zipfile.ZipFile(zip_path, 'r') as z:
        for name in z.namelist():
            entries.append((name, name.lower(), name.split('/')[-1], z.read(name)))

    try:
        os.remove(zip_path)
    except Exception:
        pass

    def process_zip_entry(entry):
        _, lower, short, data = entry
        p = {'company_info': {}, 'directors': [], 'shareholders': [],
             'auditor': {}, 'business_activity': [], 'subsidiaries': [],
             'employees': {}, 'financial': {}, 'raw_texts': []}
        try:
            # ── PDF ──────────────────────────────────────
            if lower.endswith('.pdf'):
                text = ''
                tables = []
                try:
                    with pdfplumber.open(io.BytesIO(data)) as pdf:
                        for page in pdf.pages:
                            t = page.extract_text()
                            if t:
                                text += t + '\n'
                            for table in page.extract_tables():
                                if table:
                                    tables.append(table)
                except Exception as e:
                    text = f'Error reading PDF: {str(e)}'

                if len(text.strip()) < 50 or text.startswith('Error reading PDF'):
                    ocr_text = extract_text_with_ocr(data)
                    print(f"OCR result length: {len(ocr_text)}, preview: {ocr_text[:100]}")
                    text = ocr_text.strip() or "[Scanned PDF - OCR returned no text]"

                name_upper = short.upper()
                text_upper = text[:2000].upper()
                is_aoc = 'AOC' in name_upper or 'FORM AOC' in text_upper or 'FORM NO. AOC' in text_upper
                is_mgt = 'MGT' in name_upper or 'FORM MGT' in text_upper or 'FORM NO. MGT' in text_upper or 'ANNUAL RETURN' in text_upper
                is_known_type = is_aoc or is_mgt

                if text in _NO_TEXT_PLACEHOLDERS:
                    structured = [{'heading': 'No Text Extracted', 'type': 'text',
                                    'content': 'Could not extract text from this PDF.'}]
                elif is_known_type:
                    structured = []
                else:
                    structured = extract_structured_from_text(text)
                p['raw_texts'].append({'name': short, 'text': text, 'tables': tables, 'structured': structured})

                if is_aoc:
                    extracted = extract_aoc4(text, tables)
                elif is_mgt:
                    extracted = extract_mgt7(text, tables)
                else:
                    extracted = extract_general(text)

                for k, v in extracted.get('company_info', {}).items():
                    p['company_info'].setdefault(k, v)
                for k, v in extracted.get('auditor', {}).items():
                    p['auditor'].setdefault(k, v)
                for k, v in extracted.get('employees', {}).items():
                    p['employees'].setdefault(k, v)
                p['business_activity'].extend(extracted.get('business_activity', []))
                p['subsidiaries'].extend(extracted.get('subsidiaries', []))
                p['directors'].extend(extracted.get('directors', []))

            # ── EXCEL ────────────────────────────────────
            elif lower.endswith(('.xlsx', '.xls', '.xlsm')):
                wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    rows = []
                    headers = []
                    for i, row in enumerate(ws.iter_rows(values_only=True)):
                        clean = [str(c).strip() if c is not None else '' for c in row]
                        if any(c for c in clean):
                            if i == 0:
                                headers = clean
                            else:
                                rows.append(clean)
                    if rows:
                        p['shareholders'].append({'sheet': sheet, 'headers': headers, 'rows': rows, 'source': short})

            # ── IMAGE ────────────────────────────────────
            elif lower.endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
                ext = lower.rsplit('.', 1)[-1]
                mime_map = {'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
                            'tiff': 'image/tiff', 'bmp': 'image/bmp'}
                mime_type = mime_map.get(ext, 'image/png')
                b64 = base64.b64encode(data).decode('utf-8')
                extracted = extract_data_from_image(b64, mime_type, short)
                if extracted:
                    for k, v in extracted.get('company_info', {}).items():
                        p['company_info'].setdefault(k, v)
                    for sh in extracted.get('shareholders', []):
                        p['shareholders'].append({'source': short, 'name': sh.get('name', ''), 'shares': sh.get('shares', '')})
                    for d in extracted.get('directors', []):
                        p['directors'].append([d.get('din', ''), d.get('name', ''), d.get('designation', '')])
                    for k, v in extracted.get('financial', {}).items():
                        p['financial'].setdefault(k, v)

        except Exception as e:
            print(f"Unhandled error processing {short}: {e}")
        return p

    # Process entries sequentially
    partials = []
    for entry in entries:
        partials.append(process_zip_entry(entry))

    # Merge partial results into final results
    for p in partials:
        for k, v in p['company_info'].items():
            results['company_info'].setdefault(k, v)
        for k, v in p['auditor'].items():
            results['auditor'].setdefault(k, v)
        for k, v in p['employees'].items():
            results['employees'].setdefault(k, v)
        for k, v in p['financial'].items():
            results['financial'].setdefault(k, v)
        results['directors'].extend(p['directors'])
        results['shareholders'].extend(p['shareholders'])
        results['business_activity'].extend(p['business_activity'])
        results['subsidiaries'].extend(p['subsidiaries'])
        results['raw_texts'].extend(p['raw_texts'])

    return jsonify(results)


def extract_aoc4(text, tables):
    """Extract fields from AOC-4 PDF — returns partial results dict."""
    r = {'company_info': {}, 'auditor': {}}
    patterns = {
        'CIN': [
            r'Corporate [Ii]dentity [Nn]umber[^\n]*?([A-Z]{1}[0-9]{5}[A-Z]{2}[0-9]{4}[A-Z]{3}[0-9]{6})',
            r'CIN[^\n:]*?:?\s*([A-Z]{1}[0-9]{5}[A-Z]{2}[0-9]{4}[A-Z]{3}[0-9]{6})'
        ],
        'Company Name': [
            r'Name of the [Cc]ompany[^\n]*?\n([A-Z][^\n]{3,80})',
            r'\*?Name of the company[^\n]*?\n?\s*([A-Z][A-Z\s&(),.-]{3,80}(?:LIMITED|LTD|PRIVATE|PVT)[\s.]*(?:LIMITED|LTD)?)'
        ],
        'Registered Address': [
            r'Address of the registered office[^\n]*?\n([^\n]{10,150})',
            r'\*?Address of the registered[^\n]*?\n?\s*([A-Z0-9][^\n]{10,150})'
        ],
        'Email': [
            r'e-mail[^\n]*?([a-zA-Z0-9._%+\-*]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
            r'[Ee]mail[^\n]*?([a-zA-Z0-9._%+\-*]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})'
        ],
        'Financial Year From': [
            r'[Ff]inancial year[^\n]*?[Ff]rom[^\n]*?(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'From \(DD/MM/YYYY\)[^\n]*?\n\s*(\d{2}\/\d{2}\/\d{4})'
        ],
        'Financial Year To': [
            r'[Ff]inancial year[^\n]*?[Tt]o[^\n]*?(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})',
            r'To \(DD/MM/YYYY\)[^\n]*?\n\s*(\d{2}\/\d{2}\/\d{4})'
        ],
        'Authorised Capital': [
            r'[Aa]uthorised capital[^\n]*?(\d[\d,\.]+)',
            r'[Aa]uthorised [Cc]apital of the company[^\n]*?(\d[\d,\.]+)'
        ],
        'Paid Up Capital': [
            r'[Pp]aid.?up capital[^\n]*?(\d[\d,\.]+)',
        ],
        'Turnover': [
            r'[Tt]urnover[^\n]*?(\d[\d,\.]+)',
        ],
    }

    for field, rxs in patterns.items():
        for rx in rxs:
            m = re.search(rx, text)
            if m:
                val = m.group(1).strip().replace('\n', ' ')
                if val and field not in r['company_info']:
                    r['company_info'][field] = val
                break

    auditor_patterns = {
        'Auditor Firm Name': [r'[Ff]irm [Nn]ame[^\n]*?\n\s*([A-Z][^\n]{3,80})'],
        'Auditor PAN': [r'PAN of [Aa]uditor[^\n]*?([A-Z]{5}[0-9]{4}[A-Z])'],
        'Auditor Reg No': [r'[Rr]egistration [Nn]umber[^\n]*?(\d{6,})'],
        'Signing Member': [r'[Mm]embership [Nn]umber[^\n]*?(\d{4,8})'],
    }
    for field, rxs in auditor_patterns.items():
        for rx in rxs:
            m = re.search(rx, text)
            if m:
                r['auditor'][field] = m.group(1).strip()
                break

    return r


def extract_mgt7(text, tables):
    """Extract fields from MGT-7A PDF — returns partial results dict."""
    r = {'company_info': {}, 'employees': {}, 'business_activity': [], 'subsidiaries': [], 'directors': []}
    mgt_patterns = {
        'CIN': [r'([A-Z]{1}[0-9]{5}[A-Z]{2}[0-9]{4}[A-Z]{3}[0-9]{6})'],
        'Company Name': [
            r'Name of (?:the )?[Cc]ompany[^\n]*?\n\s*([A-Z][^\n]{3,80})',
            r'([A-Z][A-Z\s&(),.-]{3,80}(?:PRIVATE\s+LIMITED|PVT\.?\s*LTD\.?))',
        ],
        'Registered Address': [
            r'[Rr]egistered office address[^\n]*?\n\s*([^\n]{10,200})',
            r'[Rr]egistered [Oo]ffice[^\n]*?\n\s*([A-Z0-9][^\n]{10,200})',
            r'(?:[Rr]egistered office address)[^\n]*?[\n:]\s*([A-Z0-9][^\n,]{5,}(?:,[^\n]{5,}){2,})',
        ],
        'Latitude': [
            r'[Ll]atitude\s+details?[^\n]*?\n\s*([\d.]+)',
            r'[Ll]atitude[^\n]*?\n\s*([\d.]+)',
            r'[Ll]atitude[^\n]*?(2[0-9]\.[\d]{4,})',
        ],
        'Longitude': [
            r'[Ll]ongitude\s+details?[^\n]*?\n\s*([\d.]+)',
            r'[Ll]ongitude[^\n]*?\n\s*([\d.]+)',
            r'[Ll]ongitude[^\n]*?(7[0-9]\.[\d]{4,})',
        ],
        'PAN': [
            r'[Pp]ermanent [Aa]ccount [Nn]umber[^\n]*?\n\s*([A-Z*]{2,5}[\d*]{4}[A-Z*])',
            r'PAN of the company[^\n]*?\n\s*([A-Z*]{5}[\d*]{4}[A-Z*])',
            r'\bPAN\b[^\n]*?\n\s*([A-Z*]{5}[\d*]{4}[A-Z*])',
        ],
        'AGM Date': [r'AGM[^\n]*?(\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4})'],
        'SRN': [r'SRN[^\n:]*?:?\s*([A-Z][0-9]{8,})'],
        'Financial Year': [r'[Ff]inancial [Yy]ear[^\n]*?(\d{4}[-–]\d{2,4})'],
    }
    for field, rxs in mgt_patterns.items():
        for rx in rxs:
            m = re.search(rx, text)
            if m:
                val = m.group(1).strip()
                if val and field not in r['company_info']:
                    r['company_info'][field] = val
                break

    emp_patterns = {
        'Total Employees': [r'[Tt]otal[^\n]*?[Ee]mployees?[^\n]*?(\d+)'],
        'Female': [r'\b[Ff]emale\b[^\n]*?(\d+)'],
        'Male': [r'\b[Mm]ale\b[^\n]*?(\d+)'],
    }
    for field, rxs in emp_patterns.items():
        for rx in rxs:
            m = re.search(rx, text)
            if m:
                r['employees'][field] = m.group(1).strip()
                break

    for table in tables:
        if not table:
            continue
        for row in table:
            if not row:
                continue
            row_text = ' '.join(str(c) for c in row if c)
            if 'Registered office address' in row_text or 'registered office address' in row_text.lower():
                for col in row[1:]:
                    val = str(col).strip() if col else ''
                    if val and len(val) > 10 and 'Registered Address' not in r['company_info']:
                        r['company_info']['Registered Address'] = val
                        break
            if 'Latitude' in row_text or 'latitude' in row_text.lower():
                for col in row[1:]:
                    val = str(col).strip() if col else ''
                    if val and re.match(r'^\d{1,2}\.\d+$', val) and 'Latitude' not in r['company_info']:
                        r['company_info']['Latitude'] = val
                        break
            if 'Longitude' in row_text or 'longitude' in row_text.lower():
                for col in row[1:]:
                    val = str(col).strip() if col else ''
                    if val and re.match(r'^\d{2,3}\.\d+$', val) and 'Longitude' not in r['company_info']:
                        r['company_info']['Longitude'] = val
                        break

    for table in tables:
        if not table:
            continue
        for row in table:
            if not row:
                continue
            row_text = ' '.join(str(c) for c in row if c)
            if any(k in row_text.upper() for k in ['NIC', 'ACTIVITY', 'TURNOVER', 'BUSINESS']):
                clean = [str(c).strip() if c else '' for c in row]
                if any(c for c in clean):
                    r['business_activity'].append(clean)

    for table in tables:
        if not table:
            continue
        for row in table:
            if not row:
                continue
            row_text = ' '.join(str(c) for c in row if c)
            if any(k in row_text.upper() for k in ['DIN', 'DIRECTOR', 'DESIGNATION']):
                clean = [str(c).strip() if c else '' for c in row]
                if any(c for c in clean):
                    r['directors'].append(clean)

    for table in tables:
        if not table:
            continue
        for row in table:
            if not row:
                continue
            row_text = ' '.join(str(c) for c in row if c)
            if any(k in row_text.upper() for k in ['SUBSIDIARY', 'ASSOCIATE', 'HOLDING']):
                clean = [str(c).strip() if c else '' for c in row]
                if any(c for c in clean):
                    r['subsidiaries'].append(clean)

    return r


def extract_general(text):
    """Extract common fields from any PDF — returns partial results dict."""
    r = {'company_info': {}}
    patterns = {
        'CIN': r'([A-Z]{1}[0-9]{5}[A-Z]{2}[0-9]{4}[A-Z]{3}[0-9]{6})',
        'Email': r'([a-zA-Z0-9._%+\-*]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,})',
        'Phone': r'(\+?91[\s\-]?[6-9]\d{9}|[6-9]\d{9})',
        'GSTIN': r'([0-9]{2}[A-Z]{5}[0-9]{4}[A-Z][1-9A-Z]Z[0-9A-Z])',
        'PAN': r'([A-Z]{5}[0-9]{4}[A-Z])',
    }
    for field, rx in patterns.items():
        m = re.search(rx, text)
        if m and field not in r['company_info']:
            r['company_info'][field] = m.group(1).strip()
    return r


@app.route('/pdf-to-word', methods=['POST'])
def pdf_to_word():
    """Convert a scanned PDF to a Word document using Google Vision OCR."""
    if not PDF2IMAGE_AVAILABLE:
        return jsonify({'error': 'pdf2image library not installed'}), 500
    if not GOOGLE_VISION_AVAILABLE:
        return jsonify({'error': 'Google Vision API not configured'}), 500

    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400

    file = request.files['file']
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Please upload a PDF file'}), 400

    pdf_data = file.read()

    try:
        images = convert_from_bytes(pdf_data, dpi=200)
    except Exception as e:
        return jsonify({'error': f'Failed to render PDF pages: {str(e)}'}), 500

    word_doc = Document()

    for i, image in enumerate(images):
        img_buffer = io.BytesIO()
        image.save(img_buffer, format='PNG')
        img_bytes = img_buffer.getvalue()
        ocr_text = _google_vision_ocr(img_bytes)
        image.close()

        heading = word_doc.add_paragraph()
        heading_run = heading.add_run(f'Page {i + 1}')
        heading_run.bold = True

        word_doc.add_paragraph(ocr_text)

    output = io.BytesIO()
    word_doc.save(output)
    output.seek(0)

    download_name = file.filename.rsplit('.', 1)[0] + '_ocr.docx'
    return send_file(
        output,
        as_attachment=True,
        download_name=download_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    app.run(debug=True)
